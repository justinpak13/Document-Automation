import docx
from datetime import datetime

def create_document():
    doc = docx.Document('Important Document.docx')

    # In this example file only 4 replacements 
    replacement_list = {
        "#Deal_Name": "new correct deal name",
        "#some_specific_word": "The correct word",
        "#manager_name": "Big Scary Manager Name",
        "#date": "Todays Date",
        "#another_specific_word":"another correct word"
    }

    # replace things in paragraphs 
    for paragraph in doc.paragraphs:
        words_to_be_replaced = set(paragraph.text.split()).intersection(set(replacement_list.keys()))
        for i in words_to_be_replaced:
            orig_text = paragraph.text
            new_text = str.replace(orig_text, i, replacement_list[i])
            paragraph.text = new_text

    #replace in tables 
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                words_to_be_replaced = set(cell.text.split()).intersection(set(replacement_list.keys()))
                for i in words_to_be_replaced:
                    orig_text = cell.text
                    new_text = str.replace(orig_text, i, replacement_list[i])
                    cell.text = new_text

    #save updated document
    time = datetime.now()
    location =f"output/Test_{time.strftime('%m_%d_%Y_%H:%M:%S')}.docx"
    doc.save(location)
    return location


if __name__ == "__main__":
    create_document()